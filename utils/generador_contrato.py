from dateutil.relativedelta import relativedelta
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from decimal import Decimal, ROUND_HALF_UP

from models import Venta
from datetime import timedelta
import tempfile
import os

# ========= Fechas en español sin locale =========
DIAS = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
MESES = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

def fecha_larga(dt):
    """Devuelve: 'miércoles, 11 de octubre de 2025' sin usar locale."""
    return f"{DIAS[dt.weekday()]}, {dt.day:02d} de {MESES[dt.month-1]} de {dt.year}"

# ========= Números a letras =========
def numero_a_letras(n: int) -> str:
    if n == 0:
        return "cero"

    unidades = ("", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve")
    especiales = {
        10: "diez", 11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
        16: "dieciséis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve"
    }
    decenas = ("", "", "veinte", "treinta", "cuarenta", "cincuenta",
               "sesenta", "setenta", "ochenta", "noventa")
    centenas = ["", "ciento", "doscientos", "trescientos", "cuatrocientos",
                "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos"]

    def _lt100(x: int) -> str:
        if x in especiales:
            return especiales[x]
        d, u = divmod(x, 10)
        if d == 0:
            return unidades[u]
        if d == 2 and u != 0:  # 21-29
            if u == 1: return "veintiuno"
            if u == 2: return "veintidós"
            if u == 3: return "veintitrés"
            if u == 6: return "veintiséis"
            return "veinti" + unidades[u]
        return f"{decenas[d]} y {unidades[u]}" if u else decenas[d]

    if n < 100:
        return _lt100(n)

    if n < 1000:
        c, r = divmod(n, 100)
        if r == 0:
            return "cien" if c == 1 else centenas[c]
        return f"{centenas[c]} {_lt100(r)}".strip()

    if n < 1_000_000:
        m, r = divmod(n, 1000)
        miles = "mil" if m == 1 else f"{numero_a_letras(m)} mil"
        return miles if r == 0 else f"{miles} {numero_a_letras(r)}"

    if n < 1_000_000_000:
        mill, r = divmod(n, 1_000_000)
        millones = "un millón" if mill == 1 else f"{numero_a_letras(mill)} millones"
        return millones if r == 0 else f"{millones} {numero_a_letras(r)}"

    return str(n)

# ========= Formateo de montos =========
def monto_formateado(valor):
    return f"$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def monto_con_letras(valor) -> str:
    # Decimal para evitar errores de flotantes
    v = Decimal(str(valor)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    entero = int(v)
    centavos = int((v - Decimal(entero)) * 100)
    return f"{numero_a_letras(entero).capitalize()} pesos con {centavos:02d}/100"

# ========= Armado de datos para plantillas =========
def preparar_datos_contrato(venta: Venta):
    cliente = venta.cliente
    garante = venta.garante

    fecha_contrato = venta.fecha
    fecha_inicio_pago = venta.fecha_inicio_pago

    frecuencia = venta.plan_pago
    if frecuencia == "diaria":
        intervalo = relativedelta(days=1)
        texto_periodicidad = "diarias"
    elif frecuencia == "semanal":
        intervalo = relativedelta(weeks=1)
        texto_periodicidad = "semanales"
    else:
        intervalo = relativedelta(months=1)
        texto_periodicidad = "mensuales"

    return {
        "cliente_nombre": f"{cliente.apellidos} {cliente.nombres}",
        "cliente_domicilio": cliente.domicilio_personal or "________",
        "cliente_localidad": cliente.localidad or "________",
        "cliente_provincia": cliente.provincia or "________",
        "cuotas": venta.num_cuotas,
        "monto_letras": f"{monto_formateado(venta.monto)} - {monto_con_letras(venta.monto)}",
        "valor_cuota_letras": f"{monto_formateado(venta.valor_cuota)} - {monto_con_letras(venta.valor_cuota)}",
        "ptf_letras": f"{monto_formateado(venta.ptf)} - {monto_con_letras(venta.ptf)}",
        "tem": f"{venta.tem:.2f}",
        "tna": f"{venta.tna:.2f}",
        "tea": f"{venta.tea:.3f}",
        "garante_nombre": f"{garante.apellidos} {garante.nombres}" if garante else "________",
        "garante_domicilio": garante.domicilio_personal if garante else "________",
        "cliente_tipo_doc": cliente.tipo_documento or "",
        "cliente_nro_doc":  cliente.nro_documento  or "",
        "garante_tipo_doc": garante.tipo_documento if garante else "________",
        "garante_nro_doc":  garante.nro_documento  if garante else "________",
        "garante_localidad": garante.localidad if garante else "________",
        "dia": fecha_contrato.day,
        "mes": MESES[fecha_contrato.month - 1],  # sin locale
        "anio": fecha_contrato.year,
        "texto_periodicidad": texto_periodicidad,
        "fecha_inicio_pago": fecha_inicio_pago.strftime("%d/%m/%Y")
    }

# ========= Generadores =========
def reemplazar_tags_doc(doc: Document, datos: dict):
    for p in doc.paragraphs:
        for key, val in datos.items():
            tag = f"{{{{{key}}}}}"
            if tag in p.text:
                p.text = p.text.replace(tag, str(val))

def _agregar_bordes_tabla(tabla) -> None:
    tblPr = tabla._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for borde in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{borde}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def insertar_tabla_vencimientos(doc: Document, venta: Venta) -> None:
    parrafo_placeholder = None
    for p in doc.paragraphs:
        if p.text.strip() == "{{vencimientos}}":
            parrafo_placeholder = p
            break
    if parrafo_placeholder is None:
        raise ValueError("No se encontró el placeholder {{vencimientos}} en la plantilla.")

    cuotas_ordenadas = sorted(venta.cuotas, key=lambda c: c.numero)
    if not cuotas_ordenadas:
        raise ValueError("La venta no tiene cuotas cargadas; no se puede generar el cronograma de vencimientos.")

    parrafo_placeholder.text = ""

    tabla = doc.add_table(rows=1, cols=3)
    try:
        tabla.style = 'Table Grid'
    except KeyError:
        _agregar_bordes_tabla(tabla)

    encabezados = ["N° Cuota", "Vencimiento", "Importe"]
    for idx, texto in enumerate(encabezados):
        celda = tabla.rows[0].cells[idx]
        celda.text = texto
        for run in celda.paragraphs[0].runs:
            run.bold = True

    for cuota in cuotas_ordenadas:
        fila = tabla.add_row()
        fila.cells[0].text = str(cuota.numero)
        fila.cells[1].text = cuota.fecha_vencimiento.strftime("%d/%m/%Y")
        fila.cells[2].text = monto_formateado(cuota.monto_original)

    anchos = (Cm(2.5), Cm(3.5), Cm(3.5))
    for fila in tabla.rows:
        for idx, ancho in enumerate(anchos):
            fila.cells[idx].width = ancho

    header_trPr = tabla.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    header_trPr.append(tblHeader)

    for fila in tabla.rows:
        trPr = fila._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)

    parrafo_placeholder._p.addnext(tabla._tbl)

def generar_contrato_word(venta: Venta, plantilla_path: str) -> str:
    salida_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    datos = preparar_datos_contrato(venta)
    doc = Document(plantilla_path)
    reemplazar_tags_doc(doc, datos)
    insertar_tabla_vencimientos(doc, venta)
    doc.save(salida_path)
    return salida_path

